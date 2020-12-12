import docx
from docx.shared import Pt
from tkfilebrowser import askopendirname
from tkinter import *
import os
import math

document = docx.Document("default.docx")
window = Tk()
window_s =Tk()
frame=Frame(window_s)
frame.pack()
frame.grid(column=0, row=24)
window.title('Autocomplete contracts')
window.geometry("800x600")
style = document.styles['Normal']
font = style.font
font.name = 'TimesNewRoman'
font.size = Pt(10)


def truncate(number, digits) -> float:
    stepper = pow(10.0, digits)
    return math.trunc(stepper * number) / stepper

def text_swipe():
    directory = askopendirname(title='Choose a directory')
    for text_par in document.paragraphs:
        text_par.text = text_par.text.replace("NUMB", numb.get())
        text_par.text = text_par.text.replace("day", day.get())
        text_par.text = text_par.text.replace("month", month.get())
        text_par.text = text_par.text.replace("year", year.get())
        text_par.text = text_par.text.replace("Prods", prods.get())
        text_par.text = text_par.text.replace("DIR", direct.get())
        text_par.text = text_par.text.replace("FIOF", fio_f.get())
        text_par.text = text_par.text.replace("default", default.get())
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for table_par in cell.paragraphs:
                        table_par.text = table_par.text.replace("DIR", direct.get())
                        table_par.text = table_par.text.replace("Prods", prods.get())
                        table_par.text = table_par.text.replace("citycomp", city_comp.get())
                        table_par.text = table_par.text.replace("UNP", UNP.get())
                        table_par.text = table_par.text.replace("unn", n_unp.get())
                        table_par.text = table_par.text.replace("RS", rs.get())
                        table_par.text = table_par.text.replace("bankname", bank_name.get())
                        table_par.text = table_par.text.replace("citybank", city_bank.get())
                        table_par.text = table_par.text.replace("BIC", BIC.get())
                        table_par.text = table_par.text.replace("tel", tel.get())
                        table_par.text = table_par.text.replace("email", email.get())
                        table_par.text = table_par.text.replace("IOF", iof.get())
    document.save(directory + "\\" + str(file_name.get()) + ".docx")
    os.startfile(directory + "\\" + str(file_name.get()) + ".docx")
    window.event_add(exit())




numb = StringVar()
day = StringVar()
month = StringVar()
year = StringVar()
prods = StringVar()
direct = StringVar()
fio_f = StringVar()
default = StringVar()
city_comp = StringVar()
UNP = StringVar()
n_unp = StringVar()
rs = StringVar()
bank_name = StringVar()
city_bank = StringVar()
BIC = StringVar()
tel = StringVar()
email = StringVar()
iof = StringVar()
file_name = StringVar()
file_name_spec = StringVar()


# NUMB
lbl = Label(window, text="Номер договора:")
lbl.grid(column=0, row=0)
numb_ent = Entry(window, textvariable=numb)
numb_ent.grid(column=1, row=0)

# DAY
lbl1 = Label(window, text="День:")
lbl1.grid(column=0, row=1)
day_ent = Entry(window, textvariable=day)
day_ent.grid(column=1, row=1)

# MONTH
lbl2 = Label(window, text="Месяц:")
lbl2.grid(column=0, row=2)
month_ent = Entry(window, textvariable=month)
month_ent.grid(column=1, row=2)

# YEAR
lbl3 = Label(window, text="Год:")
lbl3.grid(column=0, row=3)
year_ent = Entry(window, textvariable=year)
year_ent.grid(column=1, row=3)

# PRODS
lbl4 = Label(window, text="Компания:")
lbl4.grid(column=0, row=4)
prods_ent = Entry(window, textvariable=prods)
prods_ent.grid(column=1, row=4)

# DIRECT
lbl5 = Label(window, text="Должность продавца:")
lbl5.grid(column=0, row=5)
direct_ent = Entry(window, textvariable=direct)
direct_ent.grid(column=1, row=5)

# FIO_F

lbl6 = Label(window, text="ФИО продавца:")
lbl6.grid(column=0, row=6)
fio_f_ent = Entry(window, textvariable=fio_f)
fio_f_ent.grid(column=1, row=6)

# DEFAULT

lbl7 = Label(window, text="На основании чего?:")
lbl7.grid(column=0, row=7)
default_ent = Entry(window, textvariable=default)
default_ent.grid(column=1, row=7)

# CITY_COMP

lbl8 = Label(window, text="Адрес:")
lbl8.grid(column=0, row=8)
city_comp_ent = Entry(window, textvariable=city_comp)
city_comp_ent.grid(column=1, row=8)

# UNP

lbl9 = Label(window, text="Налоговая\nабривиатура(УНН,УНП,ОКПО):")
lbl9.grid(column=0, row=9)
UNP_ent = Entry(window, textvariable=UNP)
UNP_ent.grid(column=1, row=9)

# N_UNP

lbl10 = Label(window, text="Номер\nналогоплательщика:")
lbl10.grid(column=0, row=10)
n_unp_ent = Entry(window, textvariable=n_unp)
n_unp_ent.grid(column=1, row=10)

# RS

lbl11 = Label(window, text="Рассчетный счёт:")
lbl11.grid(column=0, row=11)
rs_ent = Entry(window, textvariable=rs)
rs_ent.grid(column=1, row=11)

# BANK_NAME

lbl12 = Label(window, text="Название банка:")
lbl12.grid(column=0, row=12)
bank_name_ent = Entry(window, textvariable=bank_name)
bank_name_ent.grid(column=1, row=12)

# BANK_CITY

lbl13 = Label(window, text="Адрес банка:")
lbl13.grid(column=0, row=13)
city_bank_ent = Entry(window, textvariable=city_bank)
city_bank_ent.grid(column=1, row=13)

# BIC

lbl14 = Label(window, text="BIC:")
lbl14.grid(column=0, row=14)
BIC_ent = Entry(window, textvariable=BIC)
BIC_ent.grid(column=1, row=14)

# TEL

lbl15 = Label(window, text="Телефон:")
lbl15.grid(column=0, row=15)
tel_ent = Entry(window, textvariable=tel)
tel_ent.grid(column=1, row=15)

# E-MAIL

lbl16 = Label(window, text="e-mail:")
lbl16.grid(column=0, row=16)
email_ent = Entry(window, textvariable=email)
email_ent.grid(column=1, row=16)

# IOF

lbl17 = Label(window, text="Инициалы и Фамилия:")
lbl17.grid(column=0, row=17)
iof_ent = Entry(window, textvariable=iof)
iof_ent.grid(column=1, row=17)




def windows():
    copnds = StringVar()
    summnds = StringVar()
    Pcs = StringVar()
    price = StringVar()
    summ = StringVar()
    pnds = StringVar()
    snds = StringVar()
    days = StringVar()
    months = StringVar()
    years = StringVar()
    dognum = StringVar()
    dd = StringVar()
    dm = StringVar()
    dy = StringVar()
    Cult = StringVar()
    propisnds = StringVar()
    propissummnds = StringVar()
    cop = StringVar()
    gost = StringVar()
    dost = StringVar()
    posts = StringVar()
    scomp = StringVar()
    saddres = StringVar()
    sunp = StringVar()
    sn_unp = StringVar()
    srs = StringVar()
    sbankname = StringVar()
    sbic = StringVar()
    stel = StringVar()
    se_mail = StringVar()
    sDir = StringVar()
    sIOF = StringVar()
    num = StringVar()
    window_s.title('Autocomplete contracts')
    window_s.geometry("800x600")
    def text_swipe_spec():
        directory = askopendirname(title='Choose a directory')
        for text_par_spec in document.paragraphs:
            text_par_spec.text = text_par_spec.text.replace('num', num_spec_ent.get())
            text_par_spec.text = text_par_spec.text.replace('days', days_spec_ent.get())
            text_par_spec.text = text_par_spec.text.replace('months', months_spec_ent.get())
            text_par_spec.text = text_par_spec.text.replace('years', years_spec_ent.get())
            text_par_spec.text = text_par_spec.text.replace('dogn', dognum_spec_ent.get())
            text_par_spec.text = text_par_spec.text.replace('dd', dd_spec_ent.get())
            text_par_spec.text = text_par_spec.text.replace('dm', dm_spec_ent.get())
            text_par_spec.text = text_par_spec.text.replace('dy', dy_spec_ent.get())
            text_par_spec.text = text_par_spec.text.replace('gost', gost_spec_ent.get())
            text_par_spec.text = text_par_spec.text.replace('dost', dost_spec_ent.get())
            text_par_spec.text = text_par_spec.text.replace('posts', posts_spec_ent.get())
            text_par_spec.text = text_par_spec.text.replace('propissummnds', propissummnds.get())
            text_par_spec.text = text_par_spec.text.replace('cop', cop.get())
            text_par_spec.text = text_par_spec.text.replace('porpissnds', propisnds.get())
            text_par_spec.text = text_par_spec.text.replace('copnds', copnds.get())
        for tables in document.tables:
            for rows in tables.rows:
                for cells in rows.cells:
                    for table_par_spec in cells.paragraphs:
                        price = float(price_spec_ent.get())
                        pcs = float(pcs_spec_ent.get())
                        summ = price*pcs
                        truncate(summ, 3)
                        sunds =summ*((float(pnds_spec_ent.get()) / 100) + 1)
                        truncate(sunds, 3)
                        snds = sunds-summ
                        table_par_spec.text = table_par_spec.text.replace('Cult', cult_spec_ent.get())
                        table_par_spec.text = table_par_spec.text.replace('Pcs', pcs_spec_ent.get())
                        table_par_spec.text = table_par_spec.text.replace('price', price_spec_ent.get())
                        table_par_spec.text = table_par_spec.text.replace('summ', str(pcs*price))
                        table_par_spec.text = table_par_spec.text.replace('pnds', pnds_spec_ent.get())
                        table_par_spec.text = table_par_spec.text.replace('sunds', str(sunds))
                        table_par_spec.text = table_par_spec.text.replace('snds', str(truncate(snds, 3)))
                        table_par_spec.text = table_par_spec.text.replace('scomp', scomp.get())
                        table_par_spec.text = table_par_spec.text.replace('sadres', saddres.get())
                        table_par_spec.text = table_par_spec.text.replace('sunp', sunp.get())
                        table_par_spec.text = table_par_spec.text.replace('sn_unp', sn_unp.get())
                        table_par_spec.text = table_par_spec.text.replace('srs', srs.get())
                        table_par_spec.text = table_par_spec.text.replace('sbankname', sbankname.get())
                        table_par_spec.text = table_par_spec.text.replace('sbic', sbic.get())
                        table_par_spec.text = table_par_spec.text.replace('stel', stel.get())
                        table_par_spec.text = table_par_spec.text.replace('se-mail', se_mail.get())
                        table_par_spec.text = table_par_spec.text.replace('sDir', sDir.get())
                        table_par_spec.text = table_par_spec.text.replace('sIOF', sIOF.get())
        document.save(directory + "\\" + str(file_name_spec.get()) + ".docx")
        os.startfile(directory + "\\" + str(file_name_spec.get()) + ".docx")
        window_s.event_add(exit())
    # NUM
    lbl_spec0 = Label(window_s, text="Номер спецификации:")
    lbl_spec0.grid(column=0, row=0)
    num_spec_ent = Entry(window_s, width=55)
    num_spec_ent.grid(column=1, row=0)
    #DAYS
    lbl_spec1 = Label(window_s, text="День:")
    lbl_spec1.grid(column=0, row=1)
    days_spec_ent = Entry(window_s, width=55)
    days_spec_ent.grid(column=1, row=1)
    #MONTHS
    lbl_spec2 = Label(window_s, text="Номер месяца:")
    lbl_spec2.grid(column=0, row=2)
    months_spec_ent = Entry(window_s, width=55)
    months_spec_ent.grid(column=1, row=2)
    #YEARS
    lbl_spec3 = Label(window_s, text="Год:")
    lbl_spec3.grid(column=0, row=3)
    years_spec_ent = Entry(window_s, width=55)
    years_spec_ent.grid(column=1, row=3)
    #DOGN
    lbl_spec4 = Label(window_s, text="№ Договора:")
    lbl_spec4.grid(column=0, row=4)
    dognum_spec_ent = Entry(window_s, width=55)
    dognum_spec_ent.grid(column=1, row=4)
    #DD
    lbl_spec5 = Label(window_s, text="День договора:")
    lbl_spec5.grid(column=0, row=5)
    dd_spec_ent = Entry(window_s, width=55)
    dd_spec_ent.grid(column=1, row=5)
    #DM
    lbl_spec6 = Label(window_s, text="Номер месяца договра:")
    lbl_spec6.grid(column=0, row=6)
    dm_spec_ent = Entry(window_s, width=55)
    dm_spec_ent.grid(column=1, row=6)
    #DY
    lbl_spec7 = Label(window_s, text="Год договора:")
    lbl_spec7.grid(column=0, row=7)
    dy_spec_ent = Entry(window_s, width=55)
    dy_spec_ent.grid(column=1, row=7)
    #GOST
    lbl_spec8 = Label(window_s, text="ГОСТ:")
    lbl_spec8.grid(column=0, row=8)
    gost_spec_ent = Entry(window_s, width=55)
    gost_spec_ent.grid(column=1, row=8)
    #DOST
    lbl_spec9 = Label(window_s, text="Условия поставки:")
    lbl_spec9.grid(column=0, row=9)
    dost_spec_ent = Entry(window_s, width=55)
    dost_spec_ent.grid(column=1, row=9)
    #POST
    lbl_spec10 = Label(window_s, text="Срок поставки:")
    lbl_spec10.grid(column=0, row=10)
    posts_spec_ent = Entry(window_s, width=55)
    posts_spec_ent.grid(column=1, row=10)
    #CULT
    lbl_spec11 = Label(window_s, text="Культура:")
    lbl_spec11.grid(column=0, row=11)
    cult_spec_ent = Entry(window_s, width=55)
    cult_spec_ent.grid(column=1, row=11)
    #PCS
    lbl_spec12 = Label(window_s, text="Количество:")
    lbl_spec12.grid(column=0, row=12)
    pcs_spec_ent = Entry(window_s, width=55)
    pcs_spec_ent.grid(column=1, row=12)
    #PRICE
    lbl_spec13 = Label(window_s, text="Цена:")
    lbl_spec13.grid(column=0, row=13)
    price_spec_ent = Entry(window_s, width=55)
    price_spec_ent.grid(column=1, row=13)
    #PNDS
    lbl_spec14 = Label(window_s, text="Процент НДС:")
    lbl_spec14.grid(column=0, row=14)
    pnds_spec_ent = Entry(window_s, width=55)
    pnds_spec_ent.grid(column=1, row=14)



    # #PROPISSUMNDS
    # lbl_spec11 = Label(window_s, text="Сумма с НДС прописью:")
    # lbl_spec11.grid(column=0, row=11)
    # propissummnds_spec_ent = Entry(window_s, width=55)
    # propissummnds_spec_ent.grid(column=1, row=11)
    ent_spec_btn = Button(window_s, text="Создать спецификацию", font=('calibri', 16, 'bold', 'underline'))
    ent_spec_btn.config(command=text_swipe_spec)
    ent_spec_btn.config(command=print(file_name.get()))
    ent_spec_btn.grid(column=1, row=15)


lbl17 = Label(window, text="Имя файла договора:")
lbl17.grid(column=1, row=18)
iof_ent = Entry(window, textvariable=file_name)
iof_ent.grid(column=1, row=19)
ent_btn_dog = Button(window, text="Создать договор")
ent_btn_dog.config(command= print(file_name.get()))
ent_btn_dog.config(command=text_swipe)
ent_btn_dog.grid(column=1, row=20)
def spec():
    lbl_spec1 = Label(window, text="Имя файла спецификации:")
    lbl_spec1.grid(column=1, row=23)
    file_spec_ent = Entry(window, textvariable=file_name_spec, width=55)
    file_spec_ent.grid(column=1, row=24)
    ent_btn_spec = Button(window, text="Спецификация")
    ent_btn_spec.config(command=lambda: [windows(), window.withdraw()])
    ent_btn_spec.grid(column=1, row=25)
def visable():
    ent_btn_open.grid_remove()
ent_btn_open = Button(window, text="Добавить спецификацию")
ent_btn_open.config(command=lambda: [visable(), spec()])
ent_btn_open.config()
ent_btn_open.grid(column=1, row=25)




window.mainloop()
